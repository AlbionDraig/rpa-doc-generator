import json
import logging
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)

# Default theme (fallback if template missing)
_DEFAULT_THEME = {
    "colors": {
        "primary": {"rgb": [30, 41, 59], "hex": "1E293B"},
        "secondary": {"rgb": [51, 65, 85], "hex": "334155"},
        "accent": {"rgb": [79, 70, 229], "hex": "4F46E5"},
        "dark": {"rgb": [15, 23, 42], "hex": "0F172A"},
        "table_header": {"hex": "1E293B"},
        "table_alt": {"hex": "F1F5F9"},
        "text": {"rgb": [30, 41, 59], "hex": "1E293B"},
        "muted": {"rgb": [100, 116, 139], "hex": "64748B"},
        "success": {"rgb": [5, 150, 105], "hex": "059669"},
    }
}


def _load_word_theme():
    """Load Word theme configuration from template with fallback to default."""
    theme_path = Path("app/templates/word_theme.json")
    if theme_path.exists():
        try:
            return json.loads(theme_path.read_text(encoding="utf-8"))
        except Exception as exc:
            logger.warning("Error loading Word theme template: %s, using default", exc)
            return _DEFAULT_THEME
    return _DEFAULT_THEME


def _get_theme_colors():
    """Get color theme for Word document generation."""
    theme = _load_word_theme()
    colors = theme.get("colors", _DEFAULT_THEME["colors"])
    return colors


# Brand colors — Slate / Indigo / Emerald (loaded from theme)
_theme_colors = _get_theme_colors()
COLOR_PRIMARY = RGBColor(*_theme_colors["primary"]["rgb"])
COLOR_SECONDARY = RGBColor(*_theme_colors["secondary"]["rgb"])
COLOR_ACCENT = RGBColor(*_theme_colors["accent"]["rgb"])
COLOR_DARK = RGBColor(*_theme_colors["dark"]["rgb"])
COLOR_TABLE_HEADER = _theme_colors["table_header"]["hex"]
COLOR_TABLE_ALT = _theme_colors["table_alt"]["hex"]
COLOR_TEXT = RGBColor(*_theme_colors["text"]["rgb"])
COLOR_MUTED = RGBColor(*_theme_colors["muted"]["rgb"])


def generate_sdd_word(project_data, tree, output_path, flow=None, flow_image_path=None):
    """Genera el documento SDD en formato Word (.docx) con estilo profesional."""
    try:
        doc = Document()
        _setup_document(doc)

        name = project_data.get("name", "Proyecto sin nombre")
        metadata = project_data.get("metadata", {})
        tasks = project_data.get("tasks", [])

        # Cover page
        _add_cover_page(doc, name, metadata)
        doc.add_page_break()

        # Table of contents
        _add_styled_heading(doc, "Tabla de Contenido", level=1)
        _add_toc_field(doc)
        doc.add_page_break()

        # 1. Informacion General
        _add_styled_heading(doc, "1. Informacion General", level=1)
        _add_overview(doc, project_data, flow)

        # 2. Estadisticas del Proyecto
        _add_styled_heading(doc, "2. Estadisticas del Proyecto", level=1)
        _add_stats(doc, project_data, flow)

        # 3. Flujo Principal Entre Taskbots
        _add_styled_heading(doc, "3. Flujo Principal Entre Taskbots", level=1)
        if flow_image_path and Path(flow_image_path).exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(flow_image_path), width=Inches(5.8))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(
                "Flujo principal entre taskbots: entrypoints, direccion de invocacion y contrato resumido."
            )
            run.font.size = Pt(8)
            run.font.color.rgb = COLOR_MUTED
            run.italic = True
        else:
            doc.add_paragraph("No se genero una imagen del flujo para esta ejecucion.")

        # 4. Contrato de Dependencias
        _add_styled_heading(doc, "4. Contrato de Dependencias", level=1)
        _add_dependency_contracts(doc, tasks)

        # 5. Inventario de Taskbots
        _add_styled_heading(doc, "5. Inventario de Taskbots", level=1)
        _add_task_inventory(doc, tasks)

        # 6. Contrato de Variables
        _add_styled_heading(doc, "6. Contrato de Variables", level=1)
        _add_variables_section(doc, tasks)

        # 7. Credenciales y Vaults
        _add_styled_heading(doc, "7. Credenciales y Vaults", level=1)
        _add_credentials_section(doc, project_data)

        # 8. Sistemas Externos
        _add_styled_heading(doc, "8. Sistemas Externos y Configuracion Tecnica", level=1)
        _add_systems_section(doc, project_data)

        # 9. Paquetes AA360
        _add_styled_heading(doc, "9. Paquetes AA360 Detectados", level=1)
        _add_packages_section(doc, project_data)

        # 10. Estructura del Proyecto
        _add_styled_heading(doc, "10. Estructura del Proyecto", level=1)
        _add_tree_block(doc, tree)

        # Footer divider
        _add_divider(doc)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"Documento generado automaticamente por RPA-Doc-Generator el {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}."
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        run.italic = True

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        # Force Word to update all fields (TOC) when the document is opened
        update_fields = parse_xml(
            '<w:updateFields %s w:val="true"/>' % nsdecls("w")
        )
        doc.settings.element.insert(0, update_fields)

        doc.save(str(output_file))

        logger.info("SDD Word guardado en: %s", output_file)
        return str(output_file)
    except Exception as exc:
        logger.error("Error generando SDD Word: %s", exc)
        raise


def generate_quality_word(project_data, output_path, md_content=None):
    """Genera el reporte de calidad en formato Word (.docx) con estilo profesional."""
    try:
        doc = Document()
        _setup_document(doc)

        if md_content is None:
            from app.generator.sdd_generator import _generate_quality_observations

            md_content = _generate_quality_observations(project_data)

        parsed = _parse_quality_markdown(md_content)
        project_name = parsed.get("project_name") or project_data.get("name", "Proyecto")

        # Cover
        _add_cover_page(doc, project_name, project_data.get("metadata", {}), doc_type="Calidad")
        doc.add_page_break()

        _render_quality_markdown(doc, parsed)

        # Footer
        _add_divider(doc)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Documento generado automaticamente por RPA-Doc-Generator.")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        run.italic = True

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_file))

        logger.info("Calidad Word guardado en: %s", output_file)
        return str(output_file)
    except Exception as exc:
        logger.error("Error generando Calidad Word: %s", exc)
        raise


def _parse_quality_markdown(md_content):
    lines = md_content.splitlines()
    parsed = {
        "project_name": "Proyecto",
        "sections": [],
    }

    if lines and lines[0].startswith("# "):
        title = lines[0][2:].strip()
        prefix = "Observaciones de Calidad - "
        if title.startswith(prefix):
            parsed["project_name"] = title[len(prefix):].strip() or "Proyecto"

    current_section = None
    index = 1
    while index < len(lines):
        line = lines[index]
        if line.startswith("## "):
            current_section = {"title": line[3:].strip(), "lines": []}
            parsed["sections"].append(current_section)
        elif current_section is not None:
            current_section["lines"].append(line)
        index += 1

    return parsed


def _render_quality_markdown(doc, parsed):
    for section in parsed.get("sections", []):
        title = section.get("title", "")
        lines = section.get("lines", [])
        _add_styled_heading(doc, title, level=1)

        if title == "Resumen":
            _render_quality_summary(doc, lines)
            continue

        _render_markdown_block(doc, lines)


def _render_quality_summary(doc, lines):
    items = []
    date_value = None

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        if line.lower().startswith("fecha de analisis:"):
            date_value = line.split(":", 1)[1].strip()
            continue

        match = re.match(r"^-\s+\*\*(.+?):\*\*\s*(.+)$", line)
        if match:
            items.append((match.group(1).strip(), _strip_markdown_inline(match.group(2).strip())))

    if date_value:
        items.append(("Fecha", date_value))

    if items:
        _add_info_card(doc, items)


def _render_markdown_block(doc, lines):
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()

        if not stripped:
            index += 1
            continue

        if stripped == "---":
            _add_divider(doc)
            index += 1
            continue

        if stripped.startswith("### "):
            _add_styled_heading(doc, _strip_markdown_inline(stripped[4:].strip()), level=2)
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            _render_markdown_table(doc, table_lines)
            continue

        if stripped.startswith("- "):
            style = "List Bullet 2" if raw_line.startswith("  - ") else "List Bullet"
            _add_markdown_bullet(doc, stripped[2:].strip(), style=style)
            index += 1
            continue

        text = _strip_markdown_inline(stripped)
        if text:
            doc.add_paragraph(text)
        index += 1


def _render_markdown_table(doc, table_lines):
    if len(table_lines) < 2:
        return

    headers = _split_markdown_table_row(table_lines[0])
    rows = []
    for row_line in table_lines[2:]:
        rows.append([_strip_markdown_inline(value.replace("<br>", " | ").strip()) for value in _split_markdown_table_row(row_line)])

    if headers:
        _add_table(doc, [_strip_markdown_inline(header) for header in headers], rows)


def _split_markdown_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _add_markdown_bullet(doc, text, style="List Bullet"):
    paragraph = doc.add_paragraph(style=style)
    segments = re.split(r"(\*\*.*?\*\*)", text)
    for segment in segments:
        if not segment:
            continue
        is_bold = segment.startswith("**") and segment.endswith("**")
        content = _strip_markdown_inline(segment)
        run = paragraph.add_run(content)
        run.font.size = Pt(10)
        run.bold = is_bold
        if style == "List Bullet" and (
            "⚠" in content or "hardcodeada" in content.lower() or "no tiene" in content.lower()
        ):
            run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)


def _strip_markdown_inline(text):
    cleaned = str(text)
    cleaned = re.sub(r"</?sub>", "", cleaned)
    cleaned = cleaned.replace("**", "")
    cleaned = cleaned.replace("`", "")
    return cleaned.strip()


# ---------------------------------------------------------------------------
# Helpers internos
# ---------------------------------------------------------------------------

def _setup_document(doc):
    """Configure default styles and page margins."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = COLOR_TEXT
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)


def _add_cover_page(doc, project_name, metadata, doc_type="SDD"):
    """Create a styled cover page."""
    for _ in range(3):
        doc.add_paragraph("")

    # Accent bar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500" * 40)
    run.font.color.rgb = COLOR_ACCENT
    run.font.size = Pt(8)

    doc.add_paragraph("")

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento de Diseno de Software" if doc_type == "SDD" else "Reporte de Calidad")
    run.font.size = Pt(26)
    run.font.color.rgb = COLOR_PRIMARY
    run.bold = True
    run.font.name = "Calibri Light"

    # Project name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project_name)
    run.font.size = Pt(20)
    run.font.color.rgb = COLOR_ACCENT
    run.bold = True

    # Accent bar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500" * 40)
    run.font.color.rgb = COLOR_ACCENT
    run.font.size = Pt(8)

    # Description
    description = metadata.get("description", "")
    if description:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(description)
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_MUTED
        run.italic = True

    doc.add_paragraph("")
    doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Fecha de generacion: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generado por RPA-Doc-Generator")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.italic = True


def _add_styled_heading(doc, text, level=1):
    """Add a heading with custom color styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = COLOR_PRIMARY if level == 1 else COLOR_SECONDARY
        run.font.name = "Calibri Light" if level == 1 else "Calibri"
    return heading


def _add_toc_field(doc):
    """Insert a real Word TOC field that auto-populates on open."""
    ns = nsdecls("w")
    paragraph = doc.add_paragraph()

    run = paragraph.add_run()
    run._r.append(parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % ns))

    run = paragraph.add_run()
    run._r.append(parse_xml(
        '<w:instrText %s xml:space="preserve">'
        r' TOC \o "1-3" \h \z \u '
        '</w:instrText>' % ns
    ))

    run = paragraph.add_run()
    run._r.append(parse_xml('<w:fldChar %s w:fldCharType="separate"/>' % ns))

    run = paragraph.add_run("(La tabla de contenido se generara al abrir en Word)")
    run.italic = True
    run.font.color.rgb = COLOR_MUTED
    run.font.size = Pt(9)

    run = paragraph.add_run()
    run._r.append(parse_xml('<w:fldChar %s w:fldCharType="end"/>' % ns))


def _add_divider(doc):
    """Add a visual horizontal divider."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2500" * 50)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    run.font.size = Pt(6)


def _add_bullet(doc, label, value):
    p = doc.add_paragraph(style="List Bullet")
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.color.rgb = COLOR_SECONDARY
    run_label.font.size = Pt(10)
    run_value = p.add_run(str(value))
    run_value.font.size = Pt(10)


def _add_table(doc, headers, rows):
    """Add a professionally styled table with colored header and alternating rows."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    header_row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, COLOR_TABLE_HEADER)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            if row_idx % 2 == 0:
                _set_cell_bg(cell, COLOR_TABLE_ALT)

    doc.add_paragraph("")


def _set_cell_bg(cell, hex_color):
    """Set background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_info_card(doc, items):
    """Add a styled info card with key-value pairs."""
    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, (label, value) in enumerate(items):
        label_cell = table.rows[idx].cells[0]
        label_cell.text = ""
        p = label_cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(label_cell, COLOR_TABLE_HEADER)

        value_cell = table.rows[idx].cells[1]
        value_cell.text = ""
        p = value_cell.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(10)
        run.bold = True

    doc.add_paragraph("")


def _add_observation_item(doc, text):
    """Add a styled observation bullet."""
    is_warning = "\u26a0" in text or "hardcodeada" in text.lower() or "no tiene" in text.lower()
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    if is_warning:
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)  # Red 600 for warnings


def _add_tree_block(doc, tree):
    """Add the project tree in a styled code block."""
    p = doc.add_paragraph()
    run = p.add_run(tree)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_SECONDARY


def _add_section_label(doc, text):
    """Add an inline label/badge style text."""
    p = doc.add_paragraph()
    run = p.add_run(f"  {text}  ")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_ACCENT


def _format_size(bytes_size):
    if bytes_size is None:
        return "0B"
    size = int(bytes_size)
    for unit in ["B", "KB", "MB", "GB"]:
        if size < 1024:
            return f"{size:.1f}{unit}"
        size /= 1024
    return f"{size:.1f}TB"


def _describe_error_handling(error_handling):
    states = []
    if error_handling.get("has_try"):
        states.append("try")
    if error_handling.get("has_catch"):
        states.append("catch")
    if error_handling.get("has_finally"):
        states.append("finally")
    return ", ".join(states) if states else "No explicito"


def _unique_preserve(values):
    seen = set()
    unique = []
    for value in values:
        normalized = str(value).strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        unique.append(normalized)
    return unique


# ---------------------------------------------------------------------------
# Secciones del SDD
# ---------------------------------------------------------------------------

def _add_overview(doc, project_data, flow):
    files = project_data.get("files", {})
    metadata = project_data.get("metadata", {})
    entrypoints = metadata.get("entrypoints", [])
    packages = project_data.get("packages", [])
    systems = project_data.get("systems", [])

    _add_bullet(doc, "Nombre del bot", project_data.get("name", "N/A"))
    _add_bullet(doc, "Descripcion funcional", metadata.get("description", "No disponible"))
    _add_bullet(doc, "Taskbots detectados", project_data.get("task_count", 0))
    _add_bullet(doc, "Entrypoints", ", ".join(entrypoints) if entrypoints else "No identificados")
    _add_bullet(doc, "Paquetes AA360 usados", len(packages))
    _add_bullet(doc, "Sistemas externos detectados", len(systems))
    _add_bullet(doc, "Manifest presente", "Si" if files.get("manifest_count") else "No")
    if flow:
        _add_bullet(doc, "Dependencias entre taskbots", flow.get("summary", {}).get("total_edges", 0))


def _add_stats(doc, project_data, flow):
    tasks = project_data.get("tasks", [])
    files = project_data.get("files", {})

    total_nodes = sum(t.get("node_stats", {}).get("total_nodes", 0) for t in tasks)
    decision_nodes = sum(t.get("node_stats", {}).get("decision_nodes", 0) for t in tasks)
    loop_nodes = sum(t.get("node_stats", {}).get("loop_nodes", 0) for t in tasks)
    task_calls = sum(t.get("node_stats", {}).get("task_calls", 0) for t in tasks)
    error_handlers = sum(t.get("node_stats", {}).get("error_handlers", 0) for t in tasks)
    total_size = sum(t.get("size", 0) for t in tasks)

    _add_bullet(doc, "Total de taskbots", len(tasks))
    _add_bullet(doc, "Archivos XML auxiliares", files.get("xml_count", 0))
    _add_bullet(doc, "Archivos JSON auxiliares", files.get("json_count", 0))
    _add_bullet(doc, "Nodos AA360 analizados", total_nodes)
    _add_bullet(doc, "Condiciones detectadas", decision_nodes)
    _add_bullet(doc, "Bucles o reintentos", loop_nodes)
    _add_bullet(doc, "Invocaciones runTask", task_calls)
    _add_bullet(doc, "Bloques de manejo de errores", error_handlers)
    _add_bullet(doc, "Tamano total taskbots", _format_size(total_size))
    if flow:
        _add_bullet(doc, "Aristas de flujo entre taskbots", flow.get("summary", {}).get("total_edges", 0))


def _add_dependency_contracts(doc, tasks):
    contracts = []
    for task in tasks:
        for call in task.get("task_calls", []):
            contracts.append({
                "caller": task.get("name", "Taskbot"),
                "target": call.get("target_name", "subtask"),
                "inputs": call.get("inputs", []),
                "outputs": call.get("outputs", []),
            })

    if not contracts:
        doc.add_paragraph("No se detectaron invocaciones runTask entre taskbots.")
        return

    for contract in contracts:
        _add_styled_heading(doc, f"{contract['caller']} → {contract['target']}", level=2)

        if contract["inputs"]:
            _add_section_label(doc, "Variables enviadas (entrada)")
            _add_table(
                doc,
                ["Variable", "Valor asignado"],
                [[inp.get("name", "-"), inp.get("value", "-") or "-"] for inp in contract["inputs"]],
            )
        else:
            doc.add_paragraph("Sin variables de entrada.")

        if contract["outputs"]:
            _add_section_label(doc, "Variables recibidas (salida)")
            _add_table(
                doc,
                ["Variable destino", "Variable origen"],
                [[out.get("name", "-"), out.get("value", "-") or "-"] for out in contract["outputs"]],
            )
        else:
            doc.add_paragraph("Sin variables de salida.")


def _add_task_inventory(doc, tasks):
    if not tasks:
        doc.add_paragraph("No se detectaron taskbots en el paquete exportado.")
        return

    for index, task in enumerate(tasks, start=1):
        _add_styled_heading(doc, f"{index}. {task.get('name', 'Taskbot')}", level=2)

        _add_bullet(doc, "Rol", task.get("role", "taskbot"))
        _add_bullet(doc, "Ruta", task.get("path", ""))
        _add_bullet(doc, "Entrypoint", "Si" if task.get("is_entrypoint") else "No")
        _add_bullet(doc, "Tamano", _format_size(task.get("size", 0)))

        if task.get("description"):
            _add_bullet(doc, "Descripcion declarada", task["description"])
        if task.get("developer"):
            _add_bullet(doc, "Developer declarado", task["developer"])
        if task.get("declared_date"):
            _add_bullet(doc, "Fecha declarada", task["declared_date"])

        stats = task.get("node_stats", {})
        _add_bullet(
            doc, "Resumen estructural",
            f"{stats.get('total_nodes', 0)} nodos, "
            f"{stats.get('decision_nodes', 0)} condiciones, "
            f"{stats.get('loop_nodes', 0)} bucles, "
            f"{stats.get('task_calls', 0)} llamadas a subtasks",
        )
        _add_bullet(doc, "Manejo de errores", _describe_error_handling(task.get("error_handling", {})))

        dependencies = task.get("dependencies", [])
        dep_label = ", ".join(f"{d['name']} ({d['type']})" for d in dependencies)
        _add_bullet(doc, "Dependencias", dep_label if dep_label else "Sin dependencias")

        task_calls = task.get("task_calls", [])
        if task_calls:
            summaries = [
                f"{c['target_name']} [{len(c.get('inputs', []))} in / {len(c.get('outputs', []))} out]"
                for c in task_calls
            ]
            _add_bullet(doc, "Subtasks invocadas", ", ".join(summaries))

        packages = task.get("packages", [])
        if packages:
            pkg_summary = ", ".join(f"{p['name']} {p['version']}".strip() for p in packages[:10])
            _add_bullet(doc, "Paquetes usados", pkg_summary)

        actions = _unique_preserve(task.get("actions", []))[:10]
        if actions:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run("Pasos principales:")
            run.bold = True
            run.font.color.rgb = COLOR_DARK
            for action in actions:
                doc.add_paragraph(action, style="List Bullet 2")

        comments = _unique_preserve(task.get("comments", []))[:5]
        if comments:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run("Comentarios funcionales:")
            run.bold = True
            run.font.color.rgb = COLOR_DARK
            for comment in comments:
                doc.add_paragraph(comment, style="List Bullet 2")

        systems = task.get("systems", [])
        if systems:
            sys_summary = ", ".join(f"{s['type']}: {s['value']}" for s in systems[:5])
            _add_bullet(doc, "Sistemas y endpoints", sys_summary)

        # Visual separator between taskbots
        if index < len(tasks):
            _add_divider(doc)


def _add_variables_section(doc, tasks):
    if not tasks:
        doc.add_paragraph("No se detectaron variables documentables.")
        return

    for task in tasks:
        variables = task.get("variables", {})
        input_vars = variables.get("input", [])
        output_vars = variables.get("output", [])
        internal_vars = variables.get("internal", [])

        _add_styled_heading(doc, task.get("name", "Taskbot"), level=2)
        p = doc.add_paragraph()
        run = p.add_run(
            f"Resumen: {len(input_vars)} inputs, {len(output_vars)} outputs, {len(internal_vars)} internas"
        )
        run.font.color.rgb = COLOR_MUTED
        run.italic = True

        if input_vars:
            _add_section_label(doc, "Variables de entrada")
            _add_table(
                doc,
                ["Nombre", "Tipo", "Default", "Descripcion"],
                [[v["name"], v["type"], v["default"] or "-", v["description"] or "-"] for v in input_vars],
            )

        if output_vars:
            _add_section_label(doc, "Variables de salida")
            _add_table(
                doc,
                ["Nombre", "Tipo", "Default", "Descripcion"],
                [[v["name"], v["type"], v["default"] or "-", v["description"] or "-"] for v in output_vars],
            )

        if internal_vars:
            _add_section_label(doc, "Variables internas relevantes")
            _add_table(
                doc,
                ["Nombre", "Tipo", "Scope", "Default"],
                [[v["name"], v["type"], v["scope"], v["default"] or "-"] for v in internal_vars[:12]],
            )


def _add_credentials_section(doc, project_data):
    credentials = project_data.get("credentials", [])
    if not credentials:
        doc.add_paragraph("No se detectaron credenciales o vaults en los taskbots.")
        return

    _add_table(
        doc,
        ["Credencial", "Atributo", "Vault", "Origen"],
        [
            [
                c["credential_name"],
                c.get("attribute", "-") or "-",
                c.get("vault", "-") or "-",
                c.get("source", "-"),
            ]
            for c in credentials
        ],
    )


def _add_systems_section(doc, project_data):
    systems = project_data.get("systems", [])
    if not systems:
        doc.add_paragraph("No se detectaron sistemas externos o configuraciones tecnicas relevantes.")
        return

    for system in systems:
        _add_bullet(doc, system["type"], f"{system['value']} (origen: {system['source']})")


def _add_packages_section(doc, project_data):
    packages = project_data.get("packages", [])
    if not packages:
        doc.add_paragraph("No se detectaron paquetes de AA360.")
        return

    _add_table(
        doc,
        ["Paquete", "Version"],
        [[p["name"], p["version"] or "-"] for p in packages],
    )


# ---------------------------------------------------------------------------
# Secciones ampliadas de Calidad (Priorizacion, Remediacion, Interpretacion)
# ---------------------------------------------------------------------------

def _add_quality_prioritization_table(doc, prioritization):
    """Add prioritization findings table to Word document."""
    findings = prioritization.get("priority_findings", [])
    source = prioritization.get("source", "heuristic")
    confidence = prioritization.get("confidence", "media")

    if not findings:
        doc.add_paragraph("No se detectaron hallazgos priorizables.")
        return

    # Header info
    p = doc.add_paragraph()
    run = p.add_run(f"Fuente de priorizacion: {source} | Confianza: {confidence}")
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED
    run.italic = True

    # Table: Severidad | Taskbot | Hallazgo | Por que importa
    rows = []
    for finding in findings:
        task_name = (
            finding.get("task")
            or finding.get("taskbot")
            or finding.get("task_name")
            or "General"
        )
        finding_title = (
            finding.get("title")
            or finding.get("hallazgo")
            or finding.get("finding")
            or "Hallazgo detectado"
        )
        finding_why = (
            finding.get("why")
            or finding.get("why_it_matters")
            or finding.get("reason")
            or "Sin detalle"
        )
        rows.append([
            finding.get("severity", "medio"),
            task_name,
            finding_title,
            finding_why[:100] + "..." if len(finding_why) > 100 else finding_why,
        ])

    _add_table(doc, ["Severidad", "Taskbot", "Hallazgo", "Por que importa"], rows)


def _add_quality_remediation_table(doc, prioritization):
    """Add remediation sprint plan table to Word document."""
    sprint_plan = prioritization.get("sprint_plan", [])

    if not sprint_plan:
        doc.add_paragraph("No se genero plan de remediacion.")
        return

    # Table: Prioridad | Accion | Esfuerzo | Impacto | Owner | Taskbots | Criterio de cierre
    rows = []
    for item in sprint_plan:
        tasks = item.get("tasks", [])
        done_criteria = item.get("done_criteria", [])
        tasks_value = ", ".join(tasks) if tasks else "General"
        done_criteria_value = " | ".join(done_criteria) if done_criteria else "Sin criterio"

        rows.append([
            item.get("priority", "P2"),
            item.get("action", "Accion pendiente"),
            item.get("effort", "M"),
            item.get("impact", ""),
            item.get("owner", "dev"),
            tasks_value[:50] + "..." if len(tasks_value) > 50 else tasks_value,
            done_criteria_value[:50] + "..." if len(done_criteria_value) > 50 else done_criteria_value,
        ])

    _add_table(
        doc,
        ["Prioridad", "Accion", "Esfuerzo", "Impacto", "Owner", "Taskbots", "Criterio de cierre"],
        rows
    )


def _add_quality_task_interpretations(doc, tasks, task_descriptions):
    """Add task-by-task functional interpretation to Word document."""
    taskbots = [task for task in tasks if task.get("type") == "taskbot"]

    if not taskbots:
        doc.add_paragraph("No se detectaron taskbots para interpretar.")
        return

    for task in taskbots:
        task_name = task.get("name", "Taskbot")
        description = task_descriptions.get(task_name, {})

        _add_styled_heading(doc, task_name, level=2)

        _add_bullet(doc, "Perfil AA360 sugerido", description.get("task_profile", "utilitario"))
        _add_bullet(doc, "Que hace", description.get("what_it_does", "No disponible"))
        _add_bullet(doc, "Funcion que cumple", description.get("business_function", "No disponible"))
        _add_bullet(doc, "Criticidad estimada", description.get("criticality", "media"))

        # Risks
        risks = description.get("risks", ["Sin riesgos relevantes inferidos."])
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run("Riesgos detectados:")
        run.bold = True
        run.font.color.rgb = COLOR_DARK
        for risk in risks:
            doc.add_paragraph(risk, style="List Bullet 2")

        # Recommendations
        recommendations = description.get("recommendations", ["Sin recomendaciones adicionales."])
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run("Mejoras recomendadas:")
        run.bold = True
        run.font.color.rgb = COLOR_DARK
        for recommendation in recommendations:
            doc.add_paragraph(recommendation, style="List Bullet 2")

        _add_bullet(doc, "Fuente de analisis", description.get("source", "heuristic"))
        _add_bullet(doc, "Confianza estimada", description.get("confidence", "media"))


# ---------------------------------------------------------------------------
# Observaciones de calidad (reutilizadas para Word)
# ---------------------------------------------------------------------------

def _collect_quality_observations(project_data):
    tasks = project_data.get("tasks", [])
    observations = []

    for task in tasks:
        disabled = task.get("node_stats", {}).get("disabled_nodes", 0)
        if disabled > 0:
            observations.append(
                f"{task['name']} tiene {disabled} nodo(s) deshabilitado(s). "
                "Codigo muerto puede dificultar el mantenimiento."
            )

    for task in tasks:
        eh = task.get("error_handling", {})
        if not eh.get("has_try") and task.get("type") == "taskbot":
            observations.append(
                f"{task['name']} no tiene bloques try/catch. "
                "Se recomienda manejo de errores explicito."
            )

    for task in tasks:
        eh = task.get("error_handling", {})
        if eh.get("has_try") and not eh.get("has_catch"):
            observations.append(
                f"{task['name']} tiene try pero no catch. Los errores no seran capturados."
            )

    for task in tasks:
        if not task.get("description") and task.get("type") == "taskbot":
            observations.append(f"{task['name']} no tiene descripcion declarada en cabecera.")

    for task in tasks:
        if not task.get("developer") and task.get("type") == "taskbot":
            observations.append(f"{task['name']} no tiene developer declarado en cabecera.")

    for task in tasks:
        for system in task.get("systems", []):
            value = system.get("value", "")
            if system["type"] == "file" and not value.startswith("file://$") and "$" not in value:
                observations.append(
                    f"{task['name']} usa ruta de archivo hardcodeada: {value}. "
                    "Considere usar variables globales."
                )

    credentials = project_data.get("credentials", [])
    systems = project_data.get("systems", [])
    has_db = any(s["type"] == "database" for s in systems)
    if has_db and not credentials:
        observations.append(
            "Se detectaron conexiones a base de datos pero no se encontraron "
            "credenciales via CredentialVault."
        )

    return observations
