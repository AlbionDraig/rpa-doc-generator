import base64
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from app.generator import pdf_generator, word_generator


_MIN_PNG_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO5mN2QAAAAASUVORK5CYII="
)


def _tiny_png_bytes():
    return base64.b64decode(_MIN_PNG_B64)


class _FakePisaStatus:
    def __init__(self, err=0):
        self.err = err


class ExportGeneratorsCoverageTests(unittest.TestCase):
    def _project_data(self):
        return {
            "name": "DemoBot",
            "metadata": {"description": "Demo desc", "entrypoints": ["Main"]},
            "files": {"manifest_count": 1, "xml_count": 1, "json_count": 2},
            "systems": [
                {"type": "database", "value": "jdbc:mysql://db", "source": "Main"},
                {"type": "file", "value": "C:/hardcoded/path.txt", "source": "Main"},
            ],
            "credentials": [
                {
                    "credential_name": "AA360_LOGIN",
                    "attribute": "password",
                    "vault": "Default",
                    "source": "Main",
                }
            ],
            "packages": [{"name": "Browser", "version": "1.0"}],
            "task_count": 1,
            "tasks": [
                {
                    "name": "Main",
                    "type": "taskbot",
                    "role": "main",
                    "path": "Automation Anywhere/Bots/Main",
                    "is_entrypoint": True,
                    "size": 1024,
                    "description": "Task principal",
                    "developer": "QA",
                    "declared_date": "2026-04-16",
                    "node_stats": {
                        "total_nodes": 10,
                        "decision_nodes": 1,
                        "loop_nodes": 1,
                        "task_calls": 1,
                        "error_handlers": 1,
                        "disabled_nodes": 1,
                    },
                    "error_handling": {"has_try": True, "has_catch": False, "has_finally": True},
                    "dependencies": [{"name": "Lookup", "type": "runTask"}],
                    "task_calls": [
                        {
                            "target_name": "Lookup",
                            "inputs": [{"name": "InCustomer", "value": "$Cust$"}],
                            "outputs": [{"name": "OutStatus", "value": "$Status$"}],
                        }
                    ],
                    "packages": [{"name": "Browser", "version": "1.0"}],
                    "actions": ["Abrir portal", "Consultar", "Consultar"],
                    "comments": ["Comentario A", "Comentario A", "Comentario B"],
                    "systems": [
                        {"type": "url", "value": "https://example.com"},
                        {"type": "file", "value": "C:/hardcoded/path.txt"},
                    ],
                    "variables": {
                        "input": [
                            {
                                "name": "InCustomer",
                                "type": "STRING",
                                "default": "",
                                "description": "cliente",
                            }
                        ],
                        "output": [
                            {
                                "name": "OutStatus",
                                "type": "STRING",
                                "default": "",
                                "description": "estado",
                            }
                        ],
                        "internal": [
                            {
                                "name": "Tmp",
                                "type": "STRING",
                                "scope": "local",
                                "default": "",
                            }
                        ],
                    },
                }
            ],
        }

    def test_word_generators_create_files(self):
        data = self._project_data()
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            flow_png = temp_path / "flow.png"
            flow_png.write_bytes(_tiny_png_bytes())

            sdd_docx = temp_path / "SDD_Demo.docx"
            quality_docx = temp_path / "Calidad_Demo.docx"

            sdd_result = word_generator.generate_sdd_word(
                data,
                "root\n|-- Main",
                str(sdd_docx),
                flow={"summary": {"total_edges": 1}},
                flow_image_path=str(flow_png),
            )
            quality_result = word_generator.generate_quality_word(data, str(quality_docx))

            self.assertEqual(sdd_result, str(sdd_docx))
            self.assertEqual(quality_result, str(quality_docx))
            self.assertTrue(sdd_docx.exists())
            self.assertTrue(quality_docx.exists())

    def test_word_helpers_and_observations(self):
        self.assertEqual(word_generator._format_size(None), "0B")
        self.assertEqual(word_generator._describe_error_handling({}), "No explicito")
        self.assertEqual(word_generator._describe_error_handling({"has_try": True, "has_catch": True}), "try, catch")
        self.assertEqual(word_generator._unique_preserve(["a", "a", "", " b "]), ["a", "b"])

        no_obs = word_generator._collect_quality_observations(
            {"tasks": [{"name": "Main", "type": "taskbot", "description": "ok", "developer": "dev", "systems": [], "node_stats": {}, "error_handling": {"has_try": True, "has_catch": True}}], "systems": [], "credentials": [{"credential_name": "x"}]}
        )
        self.assertEqual(no_obs, [])

        data = self._project_data()
        obs = word_generator._collect_quality_observations(data)
        self.assertTrue(any("deshabilitado" in item for item in obs))
        self.assertTrue(any("try pero no catch" in item for item in obs))
        self.assertTrue(any("hardcodeada" in item for item in obs))

    def test_word_theme_fallback_and_markdown_parsing_helpers(self):
        with patch("app.generator.word_generator.Path.exists", return_value=True), patch(
            "app.generator.word_generator.Path.read_text", side_effect=RuntimeError("broken json")
        ):
            theme = word_generator._load_word_theme()

        self.assertEqual(theme, word_generator._DEFAULT_THEME)

        parsed = word_generator._parse_quality_markdown(
            "# Observaciones de Calidad - Demo\n\n"
            "Fecha de analisis: 2026-04-16 18:40:08\n\n"
            "## Resumen\n\n"
            "- **Taskbots analizados:** 2\n"
            "- **Observaciones detectadas:** 3\n\n"
            "## Hallazgos\n\n"
            "Texto libre\n"
        )

        self.assertEqual(parsed["project_name"], "Demo")
        self.assertEqual([section["title"] for section in parsed["sections"]], ["Resumen", "Hallazgos"])
        self.assertEqual(word_generator._split_markdown_table_row("| A | B |"), ["A", "B"])
        self.assertEqual(word_generator._strip_markdown_inline("<sub>**hola** `x`</sub>"), "hola x")

    def test_word_markdown_rendering_helpers(self):
        doc = Document()
        word_generator._setup_document(doc)

        parsed = {
            "project_name": "Demo",
            "sections": [
                {
                    "title": "Resumen",
                    "lines": [
                        "Fecha de analisis: 2026-04-16 18:40:08",
                        "",
                        "- **Taskbots analizados:** 2",
                        "- **Observaciones detectadas:** 3",
                    ],
                },
                {
                    "title": "Hallazgos",
                    "lines": [
                        "- **Main** usa ruta hardcodeada",
                        "  - detalle secundario",
                        "### Subtitulo",
                        "Texto normal",
                        "| Col A | Col B |",
                        "|-------|-------|",
                        "| v1 | linea 1<br>linea 2 |",
                        "---",
                    ],
                },
            ],
        }

        word_generator._render_quality_markdown(doc, parsed)

        paragraph_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )

        self.assertIn("Resumen", paragraph_text)
        self.assertIn("Hallazgos", paragraph_text)
        self.assertIn("Subtitulo", paragraph_text)
        self.assertIn("Texto normal", paragraph_text)
        self.assertIn("ruta hardcodeada", paragraph_text.lower())
        self.assertIn("Taskbots analizados", table_text)
        self.assertIn("2026-04-16 18:40:08", table_text)
        self.assertIn("linea 1 | linea 2", table_text)

    def test_quality_word_section_helpers_cover_empty_and_populated_paths(self):
        doc = Document()
        word_generator._setup_document(doc)

        word_generator._add_quality_prioritization_table(doc, {"priority_findings": []})
        word_generator._add_quality_remediation_table(doc, {"sprint_plan": []})
        word_generator._add_quality_task_interpretations(doc, [], {})

        word_generator._add_quality_prioritization_table(
            doc,
            {
                "source": "ai",
                "confidence": "alta",
                "priority_findings": [
                    {
                        "severity": "alto",
                        "taskbot": "Main",
                        "hallazgo": "Valor hardcodeado",
                        "why_it_matters": "X" * 140,
                    }
                ],
            },
        )
        word_generator._add_quality_remediation_table(
            doc,
            {
                "sprint_plan": [
                    {
                        "priority": "P1",
                        "action": "Corregir configuracion",
                        "effort": "M",
                        "impact": "Alto",
                        "owner": "dev",
                        "tasks": ["Main", "Lookup", "OtroTaskLargo"],
                        "done_criteria": ["criterio uno", "criterio dos", "criterio tres"],
                    }
                ]
            },
        )
        word_generator._add_quality_task_interpretations(
            doc,
            [{"name": "Main", "type": "taskbot"}],
            {
                "Main": {
                    "task_profile": "principal",
                    "what_it_does": "Ejecuta flujo principal",
                    "business_function": "Coordina subtareas",
                    "criticality": "alta",
                    "risks": ["Riesgo 1"],
                    "recommendations": ["Recomendacion 1"],
                    "source": "ai",
                    "confidence": "alta",
                }
            },
        )

        paragraph_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )

        self.assertIn("No se detectaron hallazgos priorizables.", paragraph_text)
        self.assertIn("No se genero plan de remediacion.", paragraph_text)
        self.assertIn("No se detectaron taskbots para interpretar.", paragraph_text)
        self.assertIn("Fuente de priorizacion: ai | Confianza: alta", paragraph_text)
        self.assertIn("Valor hardcodeado", table_text)
        self.assertIn("Corregir configuracion", table_text)
        self.assertIn("Main", paragraph_text)
        self.assertIn("Riesgos detectados:", paragraph_text)
        self.assertIn("Recomendacion 1", paragraph_text)

    def test_pdf_helpers_and_generators(self):
        self.assertEqual(pdf_generator._escape_html("<a&b>"), "&lt;a&amp;b&gt;")
        self.assertIn("[DIR]", pdf_generator._sanitize_tree_for_pdf("📁 folder"))
        self.assertIn("<br/>", pdf_generator._fix_pre_newlines("<pre><code>a\nb</code></pre>"))
        self.assertIn("<a name=\"x\"></a>", pdf_generator._fix_heading_anchors('<h2 id="x">Title</h2>'))

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            flow_png = temp_path / "flow.png"
            flow_png.write_bytes(_tiny_png_bytes())

            sdd_pdf = temp_path / "SDD_Demo.pdf"
            quality_pdf = temp_path / "Calidad_Demo.pdf"
            captured_html = []

            def fake_create_pdf(html, dest, encoding):
                captured_html.append(html)
                dest.write(b"%PDF")
                return _FakePisaStatus(err=1)

            with patch("app.generator.pdf_generator.pisa.CreatePDF", side_effect=fake_create_pdf), patch(
                "app.generator.pdf_generator.markdown.markdown",
                return_value='<h1 id="sec">Title</h1><pre><code>a\nb</code></pre><img src="flujo_taskbots.svg"/>',
            ):
                sdd_result = pdf_generator.generate_sdd_pdf(
                    "# demo",
                    str(sdd_pdf),
                    project_name="Demo & <Bot>",
                    flow_image_path=str(flow_png),
                )
                quality_result = pdf_generator.generate_quality_pdf("# quality", str(quality_pdf), project_name="Demo")

            self.assertEqual(sdd_result, str(sdd_pdf))
            self.assertEqual(quality_result, str(quality_pdf))
            self.assertTrue(sdd_pdf.exists())
            self.assertTrue(quality_pdf.exists())
            self.assertTrue(any("data:image/png;base64" in html for html in captured_html))
            self.assertTrue(any("Reporte de Calidad" in html for html in captured_html))

    def test_pdf_generators_raise_on_error(self):
        with patch("app.generator.pdf_generator.markdown.markdown", side_effect=RuntimeError("boom")):
            with self.assertRaises(RuntimeError):
                pdf_generator.generate_sdd_pdf("# x", "out.pdf")
            with self.assertRaises(RuntimeError):
                pdf_generator.generate_quality_pdf("# x", "out.pdf")


if __name__ == "__main__":
    unittest.main()
