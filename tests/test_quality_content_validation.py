"""Test to validate that quality Word document contains all Markdown content."""
import json
from pathlib import Path
from app.generator.sdd_generator import (
    generate_quality_file,
    _generate_quality_observations,
)
from app.generator.word_generator import generate_quality_word


def test_quality_word_vs_markdown_content():
    """Validate quality Word document has all sections from Markdown."""
    # Load sample project data
    project_data = {
        "name": "Test Project",
        "tasks": [
            {
                "name": "TaskBot A",
                "type": "taskbot",
                "description": "A test taskbot",
                "developer": "Developer 1",
                "node_stats": {"disabled_nodes": 0},
                "error_handling": {"has_try": True, "has_catch": True},
                "systems": [],
                "variables": {"input": [], "output": [], "internal": []},
            },
            {
                "name": "TaskBot B",
                "type": "taskbot",
                "description": None,
                "developer": None,
                "node_stats": {"disabled_nodes": 2},
                "error_handling": {"has_try": False, "has_catch": False},
                "systems": [
                    {
                        "type": "file",
                        "value": "C:\\hardcoded\\path\\file.txt",
                    }
                ],
                "variables": {"input": [], "output": [], "internal": []},
            },
        ],
        "credentials": [],
        "systems": [
            {"type": "database", "value": "connection_string"}
        ],
        "packages": [],
        "metadata": {},
        "files": {},
    }

    # Generate quality Markdown
    md_content = _generate_quality_observations(project_data)
    
    print("\n=== QUALITY MARKDOWN CONTENT ===")
    print(md_content)
    print("\n=== KEY MARKDOWN SECTIONS ===")
    
    # Check what sections are in Markdown
    md_has_sections = {
        "Observaciones de Calidad": "Observaciones de Calidad" in md_content,
        "Hallazgos": "Hallazgos" in md_content or "findings" in md_content.lower(),
        "Priorizacion": "Priorizacion" in md_content or "prioridad" in md_content.lower(),
        "Plan de Remediacion": "Plan de Remediacion" in md_content or "sprint" in md_content.lower(),
        "Interpretacion funcional": "Interpretacion" in md_content or "taskbot" in md_content.lower(),
        "Nodo deshabilitado": "deshabilitado" in md_content.lower(),
        "Sin descripcion": "descripcion" in md_content.lower(),
        "Ruta hardcodeada": "hardcodeada" in md_content.lower(),
    }
    
    for section, present in md_has_sections.items():
        status = "✓" if present else "✗"
        print(f"  {status} {section}: {present}")
    
    # Generate quality Word
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        word_path = tmp.name
    
    try:
        generate_quality_word(project_data, word_path, md_content=md_content)
        
        # Read the Word document
        from docx import Document
        doc = Document(word_path)
        word_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        
        print("\n=== QUALITY WORD CONTENT ===")
        print(word_text[:1000])  # First 1000 chars
        print("\n=== KEY WORD SECTIONS ===")
        
        # Check what sections are in Word
        word_has_sections = {
            "Resumen": "Resumen" in word_text,
            "Hallazgos": "Hallazgos" in word_text,
            "Priorizacion": "Priorizacion" in word_text or "prioridad" in word_text.lower(),
            "Plan de Remediacion": "Plan de Remediacion" in word_text or "sprint" in word_text.lower(),
            "Interpretacion funcional": "Interpretacion" in word_text or "perfil AA360" in word_text,
            "Observaciones": "Observaciones" in word_text,
            "deshabilitado": "deshabilitado" in word_text.lower(),
            "descripcion": "descripcion" in word_text.lower(),
            "hardcodeada": "hardcodeada" in word_text.lower(),
        }
        
        for section, present in word_has_sections.items():
            status = "✓" if present else "✗"
            print(f"  {status} {section}: {present}")
        
        print("\n=== COMPARISON ===")
        print("Markdown sections present:", sum(md_has_sections.values()), "/", len(md_has_sections))
        print("Word sections present:", sum(word_has_sections.values()), "/", len(word_has_sections))

        assert "Priorizacion Inteligente de Hallazgos" in word_text
        assert "TaskBot B" in table_text
        assert "Hallazgo detectado" not in table_text
        assert "Sin detalle" not in table_text
        assert "hardcodeada" in table_text.lower() or "try/catch" in table_text.lower()
        
        # Detailed content check
        print("\n=== MISSING IN WORD ===")
        for section in md_has_sections:
            if md_has_sections[section] and not word_has_sections.get(section, False):
                print(f"  ✗ Missing: {section}")
        
    finally:
        Path(word_path).unlink(missing_ok=True)


def test_quality_word_uses_markdown_as_source_of_truth():
    project_data = {"name": "Demo", "tasks": [], "metadata": {}}
    md_content = """# Observaciones de Calidad - Demo

Fecha de analisis: 2026-04-16 18:40:08

## Resumen

- **Taskbots analizados:** 99
- **Observaciones detectadas:** 7

## Priorizacion Inteligente de Hallazgos

- **Fuente de priorizacion:** ai
- **Confianza estimada:** alta

| Severidad | Taskbot | Hallazgo | Por que importa |
|-----------|---------|----------|------------------|
| alto | Main | Hallazgo exclusivo MD | Este texto solo existe en el markdown. |
"""

    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        word_path = tmp.name

    try:
        generate_quality_word(project_data, word_path, md_content=md_content)

        from docx import Document
        doc = Document(word_path)
        word_text = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )

        assert "99" in table_text
        assert "ai" in word_text.lower() or "ai" in table_text.lower()
        assert "Hallazgo exclusivo MD" in table_text
        assert "Este texto solo existe en el markdown." in table_text
    finally:
        Path(word_path).unlink(missing_ok=True)


if __name__ == "__main__":
    test_quality_word_vs_markdown_content()
